"""
build_docs.py — Generate judge-qa.pdf and judge-qa.docx from content
Run: python build_docs.py
"""

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, KeepTogether, PageBreak
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ── Colour palette ────────────────────────────────────────────────────────────
INDIGO   = colors.HexColor('#4F46E5')
VIOLET   = colors.HexColor('#7C3AED')
EMERALD  = colors.HexColor('#059669')
AMBER    = colors.HexColor('#D97706')
CYAN     = colors.HexColor('#0891B2')
RED      = colors.HexColor('#DC2626')
SLATE900 = colors.HexColor('#0F172A')
SLATE800 = colors.HexColor('#1E293B')
SLATE700 = colors.HexColor('#334155')
SLATE300 = colors.HexColor('#CBD5E1')
SLATE200 = colors.HexColor('#E2E8F0')
WHITE    = colors.white
LIGHT_BG = colors.HexColor('#F8FAFC')
CARD_BG  = colors.HexColor('#F1F5F9')
BORDER   = colors.HexColor('#E2E8F0')

# ── Styles ────────────────────────────────────────────────────────────────────
styles = getSampleStyleSheet()

def S(name, **kw):
    return ParagraphStyle(name, **kw)

COVER_TITLE = S('CoverTitle', fontSize=36, leading=44, alignment=TA_CENTER,
                textColor=SLATE900, fontName='Helvetica-Bold', spaceAfter=6)
COVER_SUB   = S('CoverSub',   fontSize=14, leading=20, alignment=TA_CENTER,
                textColor=colors.HexColor('#475569'), spaceAfter=4)
COVER_META  = S('CoverMeta',  fontSize=11, leading=16, alignment=TA_CENTER,
                textColor=colors.HexColor('#64748B'))

SEC_HEAD    = S('SectionHead', fontSize=16, leading=22, fontName='Helvetica-Bold',
                textColor=SLATE900, spaceBefore=18, spaceAfter=8)
SEC_SUB     = S('SecSub',      fontSize=11, leading=15, fontName='Helvetica',
                textColor=colors.HexColor('#475569'), spaceBefore=2, spaceAfter=10)

Q_NUM       = S('QNum',   fontSize=9,  leading=13, fontName='Helvetica-Bold',
                textColor=WHITE, backColor=INDIGO, borderPadding=(2,6,2,6))
QUESTION    = S('Question', fontSize=13, leading=19, fontName='Helvetica-Bold',
                textColor=SLATE900, spaceBefore=4, spaceAfter=4)
ANSWER      = S('Answer',   fontSize=11, leading=17, fontName='Helvetica',
                textColor=colors.HexColor('#374151'), spaceAfter=6)
CATEGORY    = S('Category', fontSize=9,  leading=13, fontName='Helvetica-Bold',
                textColor=INDIGO, spaceBefore=6, spaceAfter=2)
BULLET      = S('Bullet',   fontSize=10.5, leading=16, fontName='Helvetica',
                textColor=colors.HexColor('#374151'), leftIndent=16, spaceAfter=2)
SCREEN_TXT  = S('Screen',   fontSize=9,  leading=14, fontName='Helvetica-Oblique',
                textColor=colors.HexColor('#6366F1'), leftIndent=16, spaceAfter=4)
STEP_HEAD   = S('StepHead', fontSize=12, leading=18, fontName='Helvetica-Bold',
                textColor=SLATE900, spaceBefore=4, spaceAfter=2)
STEP_BODY   = S('StepBody', fontSize=10.5, leading=16, fontName='Helvetica',
                textColor=colors.HexColor('#374151'), leftIndent=20, spaceAfter=4)

TABLE_HEADER_STYLE = S('THead', fontSize=9.5, fontName='Helvetica-Bold',
                        textColor=WHITE, alignment=TA_LEFT)
TABLE_CELL_STYLE   = S('TCell', fontSize=9.5, fontName='Helvetica',
                        textColor=SLATE900, alignment=TA_LEFT)

# ── Helpers ───────────────────────────────────────────────────────────────────
def HR():
    return HRFlowable(width='100%', thickness=1, color=BORDER, spaceAfter=12, spaceBefore=12)

def card_table(content_rows):
    """Wrap content in a light-grey card."""
    t = Table([[c] for c in content_rows], colWidths=['100%'])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), CARD_BG),
        ('BOX',        (0,0), (-1,-1), 0.5, BORDER),
        ('ROWBACKGROUNDS', (0,0), (-1,-1), [CARD_BG]),
        ('TOPPADDING',  (0,0), (-1,-1), 8),
        ('BOTTOMPADDING',(0,0), (-1,-1), 6),
        ('LEFTPADDING', (0,0), (-1,-1), 12),
        ('RIGHTPADDING',(0,0), (-1,-1), 12),
    ]))
    return t

def color_hex(c):
    """Convert a ReportLab Color to a 6-digit hex string like #4f46e5."""
    r = int(round(c.red   * 255))
    g = int(round(c.green * 255))
    b = int(round(c.blue  * 255))
    return f'#{r:02x}{g:02x}{b:02x}'

def category_badge(text, color=INDIGO):
    return Paragraph(f'<font color="{color_hex(color)}"><b>{text}</b></font>', CATEGORY)

def qa_block(num, cat, question, answer_lines, screen=None, proofs=None, color=INDIGO):
    items = []
    items.append(category_badge(cat, color))
    items.append(Paragraph(f'<b>Q{num}. {question}</b>', QUESTION))
    for line in answer_lines:
        items.append(Paragraph(line, ANSWER))
    if screen:
        items.append(Paragraph(f'📺 {screen}', SCREEN_TXT))
    if proofs:
        for p in proofs:
            items.append(Paragraph(f'  • {p}', BULLET))
    return KeepTogether([card_table(items), Spacer(1, 8)])

# ── Content ───────────────────────────────────────────────────────────────────
def build_pdf(path):
    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=22*mm, bottomMargin=22*mm,
        title='SKILL FORGE — Judge Q&A + Demo Guide',
        author='SkillForge AI',
    )

    story = []
    W = A4[0] - 40*mm   # usable width

    # ── COVER ─────────────────────────────────────────────────────────────────
    story.append(Spacer(1, 20*mm))
    story.append(Paragraph('SKILL FORGE', COVER_TITLE))
    story.append(Paragraph('Autonomous Multi-Agent Career Intelligence Platform', COVER_SUB))
    story.append(Paragraph('Judge Q&amp;A Reference + Live Demo Guide', COVER_META))
    story.append(Spacer(1, 10*mm))

    # Stats table
    stats = Table(
        [['50+', '9', '7', '100%'],
         ['Judge Q&As', 'Agent Types', 'Demo Steps', 'Requirements Met']],
        colWidths=[W/4]*4,
    )
    stats.setStyle(TableStyle([
        ('ALIGN',       (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME',    (0,0), (-1,0),  'Helvetica-Bold'),
        ('FONTSIZE',    (0,0), (-1,0),  22),
        ('TEXTCOLOR',   (0,0), (-1,0),  INDIGO),
        ('FONTNAME',    (0,1), (-1,1),  'Helvetica'),
        ('FONTSIZE',    (0,1), (-1,1),  9),
        ('TEXTCOLOR',   (0,1), (-1,1),  colors.HexColor('#64748B')),
        ('BACKGROUND',  (0,0), (-1,-1), CARD_BG),
        ('BOX',         (0,0), (-1,-1), 0.5, BORDER),
        ('INNERGRID',   (0,0), (-1,-1), 0.3, BORDER),
        ('TOPPADDING',  (0,0), (-1,-1), 8),
        ('BOTTOMPADDING',(0,0),(-1,-1), 8),
    ]))
    story.append(stats)
    story.append(Spacer(1, 8*mm))
    story.append(HR())

    # ── PART 1: DEMO GUIDE ────────────────────────────────────────────────────
    story.append(Paragraph('🚀 Part 1 — Live Demo Guide (Run This First)', SEC_HEAD))
    story.append(Paragraph('Navigate to http://localhost:5173/demo · Click "Run Live Autonomous Career Analysis"', SEC_SUB))

    alert = Table([[Paragraph(
        '<b>START HERE:</b> Go to /demo → Select a goal → Click the run button. '
        'Watch 7 agents activate with real API calls. Takes ~15 seconds. '
        'After completion, navigate: Career Twin → Simulation Lab → Explainability Console.',
        S('Alert', fontSize=10.5, leading=16, fontName='Helvetica', textColor=colors.HexColor('#065F46')))
    ]], colWidths=[W])
    alert.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#ECFDF5')),
        ('BOX',        (0,0), (-1,-1), 1, EMERALD),
        ('TOPPADDING', (0,0), (-1,-1), 10),
        ('BOTTOMPADDING',(0,0),(-1,-1), 10),
        ('LEFTPADDING',(0,0),(-1,-1), 12),
        ('RIGHTPADDING',(0,0),(-1,-1), 12),
    ]))
    story.append(alert)
    story.append(Spacer(1, 8))

    demo_steps = [
        ('1', 'Open Landing Page (http://localhost:5173)',
         'Show hero with 7-agent badge, "Live Demo" button (top-right), Frontier Modules grid.',
         'Screen: Landing — hero, agent terminal animation, frontier modules'),
        ('2', 'Click "🚀 Live Demo" → Select "Medical Doctor"',
         'Professional domain (not generic tech) — proves domain-specific intelligence.',
         'Screen: /demo — goal selection, 5 career options'),
        ('3', 'Click "Run Live Autonomous Career Analysis"',
         'GoalAgent → DecomposeAgent → DiagnosticAgent → ScoringAgent → CurriculumAgent → MarketAgent → SimulationAgent. Real API calls, not mocked. Visible in SSE stream.',
         'Screen: /demo — agent pills glowing, execution log building live'),
        ('4', 'Review the Completion Summary',
         'Skills mapped, sessions planned, market demand, open jobs, projected salary, agents used.',
         'Screen: /demo — completion card with 6-metric grid'),
        ('5', 'Click "Career Digital Twin"',
         'Competency radar, skill mastery bars, 12-month trajectory, live market intelligence panel.',
         'Screen: /career-twin — radar chart, opportunity radar, salary forecast'),
        ('6', 'Simulation Lab → Type "Machine Learning" → Simulate',
         '+30% salary, demand jump from 58→88, faster hiring, Agent Verdict with confidence.',
         'Screen: /simulation — 4 metric cards, timeline, agent verdict banner'),
        ('7', 'Open Explainability Console',
         'Every agent decision logged with reasoning, timestamps, confidence bars, debate transcripts.',
         'Screen: /explain — agent map, decision cards, debate history'),
    ]
    for num, title, body, screen in demo_steps:
        step_rows = [
            Paragraph(f'<b>Step {num}: {title}</b>', STEP_HEAD),
            Paragraph(body, STEP_BODY),
            Paragraph(f'📺 {screen}', SCREEN_TXT),
        ]
        story.append(KeepTogether([card_table(step_rows), Spacer(1, 6)]))

    story.append(HR())

    # ── PART 2: DIFFERENTIATOR TABLE ─────────────────────────────────────────
    story.append(Paragraph('⚡ Part 2 — Key Differentiators vs ChatGPT/Claude', SEC_HEAD))

    headers = ['Feature', 'ChatGPT / Claude', 'SKILL FORGE']
    rows = [
        ['Architecture',     'Single LLM, prompt-response',           '9-agent autonomous pipeline'],
        ['Memory',           'Session-only, resets',                  'Persistent career memory (JSON → DB)'],
        ['Planning',         'Suggests a plan in chat',               'Executes real 14-21 day personalized plan'],
        ['Adaptation',       'No feedback loop',                      'AdaptorAgent + SkillDriftAgent modify plan live'],
        ['Agent Debate',     'N/A',                                   '3 agents debate before every plan change'],
        ['Simulation',       'Text description only',                 'SimulationAgent projects salary/demand/time'],
        ['Market Data',      'Training data cutoff',                  'MarketAgent: live demand scores + job counts'],
        ['Digital Twin',     'None',                                  'Radar + trajectory + opportunity radar'],
        ['Explainability',   'Opaque',                                'Full reasoning chain, every decision logged'],
        ['Domain Specificity','Generic answers',                       'Real MCQs: STEMI, IPC, M25 concrete'],
        ['LLM Resilience',   'Single provider',                       'Gemini → Groq auto-fallback on 429'],
    ]

    diff_data = [[Paragraph(h, TABLE_HEADER_STYLE) for h in headers]]
    for r in rows:
        diff_data.append([Paragraph(r[0], TABLE_CELL_STYLE),
                          Paragraph(r[1], TABLE_CELL_STYLE),
                          Paragraph('<font color="#059669"><b>✓ ' + r[2] + '</b></font>', TABLE_CELL_STYLE)])

    diff_table = Table(diff_data, colWidths=[W*0.22, W*0.33, W*0.45])
    diff_table.setStyle(TableStyle([
        ('BACKGROUND',   (0,0), (-1,0),  SLATE900),
        ('ROWBACKGROUNDS',(0,1),(-1,-1), [WHITE, CARD_BG]),
        ('BOX',          (0,0), (-1,-1), 0.5, BORDER),
        ('INNERGRID',    (0,0), (-1,-1), 0.3, BORDER),
        ('TOPPADDING',   (0,0), (-1,-1), 6),
        ('BOTTOMPADDING',(0,0), (-1,-1), 6),
        ('LEFTPADDING',  (0,0), (-1,-1), 8),
        ('RIGHTPADDING', (0,0), (-1,-1), 8),
        ('VALIGN',       (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(diff_table)
    story.append(HR())

    # ── Q&A SECTIONS ─────────────────────────────────────────────────────────
    qa_sections = [
        ('🔧 Part 3 — Technical Questions', 'Technical questions covering stack, LLM fallback, evaluation, and session design.', [
            (1,  'Technical', 'What is the tech stack?',
             ['<b>Frontend:</b> React 18 + Vite + TailwindCSS — fast, modern, responsive.',
              '<b>Backend:</b> Node.js (ES Modules) + Express.',
              '<b>LLMs:</b> Gemini 2.0 Flash (primary) + Groq llama-3.3-70b (free fallback).',
              '<b>Storage:</b> JSON file store — upgrade-ready to PostgreSQL.'],
             None, ['server/index.js', 'client/vite.config.js']),
            (2,  'Technical', 'What happens if Gemini quota is exceeded?',
             ['GeminiService.js has <b>automatic Groq fallback</b>. HTTP 429 immediately switches to Groq (free, llama-3.3-70b). All callers use the same service — transparent. If both fail, rule-based logic takes over. <b>System never crashes.</b>'],
             None, ['server/services/GeminiService.js', 'Tested with real 429 error']),
            (3,  'Technical', 'Are diagnostic questions generic or domain-specific?',
             ['Three-tier system: (1) RuleBase — real MCQs for Medicine (STEMI/troponin), Law (IPC), Civil (M25 concrete). (2) Gemini — domain-specific for any other field. (3) Knowledge-testing fallback templates. Questions test actual professional knowledge — never "why is X important?"'],
             'Enter "I want to be a Doctor" — see STEMI question', ['server/agent/RuleBase.js', 'server/agent/QuizGenerator.js']),
            (4,  'Technical', 'How does session evaluation work?',
             ['Evaluator.js uses 3-tier scoring: (1) Gemini AI — analyzes response against criteria, provides strengths/weaknesses. (2) Keyword + depth scoring. (3) Normalized A/B/C/D/F grade with 0-100 score.'],
             'Complete a session, show AI feedback', ['server/agent/Evaluator.js']),
            (5,  'Technical', 'How is the learning plan personalized?',
             ['PlanBuilder applies proficiency multiplier: skills &lt;60% get more days, &gt;80% get compressed. Uses learner profile (beginner/intermediate/advanced) from goal analysis + profiling questionnaire. Sorted weakest-first. AdaptorAgent re-evaluates after every 3 sessions.'],
             'Dashboard — Plan tab, see day allocation', ['server/agent/PlanBuilder.js', 'server/agent/Adaptor.js']),
            (6,  'Technical', 'What data does the system store?',
             ['Each user = one UUID-named JSON file in server/data/. Contains: goal, skillTree, diagnosticQuestions, scores, learningPlan, all sessions, adaptations, agentDecisions log, agentDebates, report. Upgrade to PostgreSQL = swap 2 methods in SmartAgent.js.'],
             None, ['server/data/*.json', 'server/agent/SmartAgent.js lines 406-414']),
            (7,  'Technical', 'How does the What-If simulator work?',
             ['SimulationAgent sends current goal/skills/domain + proposed skill to Gemini/Groq. Gets: salary delta, demand shift, hiring timeline improvement, learning timeline, alternatives, risk factors. Falls back to SKILL_BOOST_MAP + DOMAIN_SALARY_MAP if LLM unavailable.'],
             '/simulation → type "Kubernetes" → simulate', ['server/agent/SimulationAgent.js']),
            (8,  'Technical', 'How does the Live Demo mode work technically?',
             ['demo.js uses <b>Server-Sent Events (SSE)</b> — real-time server→browser stream. No WebSocket overhead. Server runs real API calls, emits each step as SSE event. React DemoMode.jsx listens via EventSource and animates agents in real-time.'],
             None, ['server/routes/demo.js', 'client/src/pages/DemoMode.jsx — real API calls']),
            (9,  'Technical', 'How do you prevent React UI from flickering/rolling back?',
             ['Session.jsx uses <b>cancelled flag cleanup pattern</b>: let cancelled = false, set to true on cleanup. API callbacks check !cancelled before setState. Prevents React 18 Strict Mode double-invoke from resetting phase after user has progressed. Real bug we diagnosed and fixed.'],
             None, ['client/src/pages/Session.jsx — useEffect cleanup']),
            (10, 'Technical', 'What is the latency for full goal analysis?',
             ['First analysis: 3-8s with Gemini, &lt;1s rule-based. Session evaluation: 2-5s LLM, instant rule-based. All calls have 30s AbortSignal timeout. 60s client timeout handles edge cases.'],
             None, ['All API calls include AbortSignal.timeout(30000)']),
            (11, 'Technical', 'How does confidence calibration work?',
             ['User declares confidence before each session. Report compares declared vs actual. Consistently saying "Good" but scoring &lt;60% = overconfidence bias. Saying "Poor" but scoring &gt;75% = underconfidence. Both patterns highlighted in report.'],
             None, ['client/src/pages/Session.jsx — confidence phase']),
            (12, 'Technical', 'What does the Skill Drift Agent detect?',
             ['_detectSkillDrift() compares avg of first 2 sessions on a skill vs last 2. Drop &gt;20pts = drift logged: "Performance dropped Xpts — early mastery superficial. Recommending spaced repetition." Real learning science (Ebbinghaus forgetting curve).'],
             '/explain — filter AdaptorAgent decisions', ['SmartAgent.js — _detectSkillDrift()']),
            (13, 'Technical', 'Why SSE instead of WebSockets for demo?',
             ['SSE is simpler, HTTP/1.1 compatible, auto-reconnects, works through proxies without config. For unidirectional server→client stream (all we need), SSE is the right tool. No bidirectional complexity overhead.'],
             None, ['Correct engineering choice for the use case']),
            (14, 'Technical', 'How is the Learn-First session flow designed?',
             ['Before every challenge: concept definition, 4 domain-specific points, real example, pro tip. Uses conceptSummary from ChallengeEngine. Flow: Confidence → Learn → Warmup → Challenge. Primes concept before testing — spaced learning theory.'],
             'Navigate to a session day', ['Session.jsx LearnPhase', 'ChallengeEngine.js conceptSummary']),
            (15, 'Technical', 'What happens if both Gemini and Groq are down?',
             ['Three-layer resilience: LLM → Rule-based → Static fallback. RuleBase.js has real quiz banks for major domains. QuizGenerator has 6 knowledge-testing fallback templates. ChallengeEngine has buildDynamicChallenge() fallback. System is fully functional offline for supported domains.'],
             None, ['server/agent/RuleBase.js', 'server/agent/QuizGenerator.js fallback templates']),
        ]),

        ('🏗️ Part 4 — Architecture Questions', 'System design, state management, agent pipeline, and scalability architecture.', [
            (16, 'Architecture', 'Walk me through the full agent pipeline.',
             ['1. GoalAgent (SmartAgent.processGoal) — parses goal, extracts domain, builds learner profile.',
              '2. DecomposeAgent (SkillDecomposer) — decomposes into 4-6 ordered skills with topics/time.',
              '3. DiagnosticAgent (QuizGenerator) — generates 5 targeted MCQs.',
              '4. ScoringAgent (Evaluator.scoreDiagnostic) — scores diagnostic, identifies gaps.',
              '5. CurriculumAgent (PlanBuilder) — builds personalized day plan from gaps.',
              '6. EvaluatorAgent (Evaluator.scoreSession) — scores each practice session.',
              '7. AdaptorAgent (Adaptor + AgentDebate) — adapts plan, runs internal debate.',
              '8. MarketAgent — real-time market intelligence.',
              '9. SimulationAgent — what-if career projection.'],
             '/demo — watch all 7 activate in sequence', ['/explain — agent activity map']),
            (17, 'Architecture', 'How is state shared between agents?',
             ['All agents receive the session object as shared state carrier. SmartAgent.js is the orchestrator — passes session to each agent, merges results. No direct agent-to-agent dependencies. This is the <b>Blackboard pattern</b> — agents read/write to central state through orchestrator.'],
             None, ['SmartAgent.js — session object threading']),
            (18, 'Architecture', 'Is the system truly agentic or just prompt chaining?',
             ['Key distinctions: (1) Independent decision logic — AdaptorAgent decides based on data, not LLM output. (2) Agent Debate — 3 agents with different biases argue before decisions. (3) SkillDriftAgent runs continuously and self-initiates alerts. (4) Re-routing — rule-base bypasses LLM entirely when available.'],
             None, ['Adaptor.js independent logic', 'AgentDebate.js']),
            (19, 'Architecture', 'How would you scale to millions of users?',
             ['Storage: PostgreSQL (JSONB + UUID index). LLM tier: BullMQ/Redis queue with per-user quotas. Agent parallelism: Steps 3+market run concurrently via Promise.all — already async-ready. Caching: Redis for domain skill trees (24h TTL). Horizontal: Express is stateless, loadbalancer-ready.'],
             None, []),
            (20, 'Architecture', 'What is the memory model — short-term vs long-term?',
             ['Short-term: active session object — current day, challenge, confidence, recent scores. Long-term: persistent session file — complete history of goals, diagnostics, sessions, adaptations, debates, mastery trajectory. Survives server restarts. Enables cross-week personalization.'],
             None, ['server/data/[userId].json']),
            (21, 'Architecture', 'How is the frontend architecture structured?',
             ['React 18 + lazy loading (code-split per page). 10 routes. Shared state via localStorage (userId). All comms through api.js (single source). No Redux — simple enough without it. TailwindCSS for zero-overhead styling.'],
             None, ['client/src/App.jsx', 'client/src/utils/api.js']),
            (22, 'Architecture', 'Why not use LangChain?',
             ['Deliberate choice. LangChain adds abstraction overhead + version fragility. Our agents are well-defined with custom logic (rule-based routing, domain knowledge banks, Agent Debate). Hand-rolling gives full control over error handling and fallback logic — critical for reliability.'],
             None, []),
            (23, 'Architecture', 'How does the agent architecture support new agents?',
             ['Each agent is a standalone class with no inter-agent dependencies. Adding a new agent: (1) Create class in server/agent/. (2) Import in SmartAgent.js. (3) Call at appropriate pipeline step. (4) Add decisions to agentDecisions array. <b>Open/Closed Principle</b> — no existing code changes required.'],
             None, []),
        ]),

        ('🤖 Part 5 — Agentic Reasoning Questions', 'Agent intelligence, debate engine, human-in-the-loop, and explainability.', [
            (24, 'Agentic', 'What makes your system agentic vs a simple chatbot?',
             ['(1) Plans independently — DecomposeAgent decides skill ordering. (2) Takes action — AdaptorAgent inserts/removes days. (3) Self-corrects — SkillDriftAgent detects drops and logs intervention. (4) Debates internally — 3 agent voices before adaptations. (5) Persists across time — memory informs future decisions.'],
             None, []),
            (25, 'Agentic', 'Explain the Agent Debate mechanism.',
             ['When adaptation is triggered, 3 agents with different biases evaluate the same data:',
              '<b>AdvocateAgent</b> (optimistic) — trusts learner, argues continuation.',
              '<b>CriticAgent</b> (cautious) — flags risk, argues for review.',
              '<b>AnalystAgent</b> (data-driven) — purely quantitative, highest confidence weight.',
              'Each votes with confidence weight. Weighted majority wins. Visible in Explainability Console.'],
             '/explain — Agent Debate History section', ['server/agent/AgentDebate.js']),
            (26, 'Agentic', 'Does the system do multi-step reasoning?',
             ['Full chain: Goal → Domain Detection → Skill Mapping → Dependency Analysis → Knowledge Gap Scoring → Time Allocation → Session Sequencing → Challenge Generation → Evaluation → Adaptation → Drift Detection. Each step logged. Expandable in Explainability Console.'],
             '/explain — expand a ScoringAgent decision', []),
            (27, 'Agentic', 'Is there Human-in-the-Loop? Where?',
             ['Yes, at multiple points: (1) Confidence self-rating before sessions. (2) Profiling questionnaire adjusts plan. (3) Diagnostic quiz answers drive gap analysis. (4) Override: user can navigate to any session day. (5) Simulation Lab: user actively queries what-if scenarios.'],
             None, []),
            (28, 'Agentic', 'What is the Career Digital Twin concept?',
             ['A Digital Twin is a virtual model that updates in real-time. Ours models: competency radar (current vs target per skill), mastery trajectory (month-by-month readiness), market position (skills vs job requirements), opportunity gaps. Evolves after every session.'],
             '/career-twin — radar chart + opportunity radar', ['server/agent/SimulationAgent.js forecastTrajectory()']),
            (29, 'Agentic', 'How does the What-If Simulator change career decisions?',
             ['Traditional: vague, uninformed choices. Ours: "Should I learn Kubernetes or React?" → +22% salary vs +15%, specific demand scores, hiring time improvement. Agent Verdict adds confidence. Turns vague choice into <b>data-driven investment decision</b> before spending 10 weeks learning.'],
             '/simulation — type any skill, see projected impact', []),
            (30, 'Agentic', 'How does Skill Drift Detection address learning science?',
             ['Ebbinghaus forgetting curve (1885): skills not reinforced decay rapidly. SkillDriftAgent detects algorithmically: early session scores vs recent scores on same skill. Drop &gt;20pts = superficial memorization detected. Logs intervention, recommends spaced repetition. <b>Evidence-based adaptive learning.</b>'],
             None, ['SmartAgent.js _detectSkillDrift()']),
            (31, 'Agentic', 'What is the "Learn-First" design innovation?',
             ['Traditional systems test before teaching. We invert: concept summary (definition + 4 points + example + pro tip) before every challenge. Grounded in spaced learning theory — priming reduces cognitive load, improves retention. Confidence rating calibrates metacognitive accuracy.'],
             'Navigate to a session day', ['Session.jsx LearnPhase', 'ChallengeEngine.js conceptSummary']),
        ]),

        ('🔒 Part 6 — Security Questions', 'API key security, input validation, XSS/CSRF, auth roadmap.', [
            (32, 'Security', 'How are API keys secured?',
             ['Keys stored in .env (gitignored), read at runtime via process.env. GeminiService reads fresh per call via getter — never cached in memory. Never committed (verified in .gitignore). Production: AWS/GCP Secrets Manager or HashiCorp Vault.'],
             None, ['server/services/GeminiService.js get apiKey()', '.env in .gitignore']),
            (33, 'Security', 'What input validation is in place?',
             ['Goal text: length validated (5-500 chars). Session submission: all required fields validated before LLM call. Express JSON body limited to 2MB. Consistent {success, data, error} envelope — no raw stack traces exposed.'],
             None, ['server/routes/goal.js', 'server/routes/session.js']),
            (34, 'Security', 'Is there XSS/CSRF protection?',
             ["React's JSX inherently escapes all user-provided content (no dangerouslySetInnerHTML). CSRF not applicable — no cookies used (userId in localStorage). CORS configured to allow only localhost:5173/5174 in dev."],
             None, []),
            (35, 'Security', 'How would you handle auth in production?',
             ['JWT-based auth with refresh token rotation. Auth middleware on all routes. userId extracted from verified JWT (not user-supplied). Bcrypt password hashing or OAuth2 via Google/GitHub. Session data encrypted at rest (AES-256). Auth is additive — no architectural change needed.'],
             None, []),
            (36, 'Security', 'Is there rate limiting?',
             ['Not in hackathon version — unnecessary for demo. Production: express-rate-limit with per-IP + per-userId limits. LLM calls with per-user quotas. 60s client timeout + 30s server AbortSignal prevent runaway requests.'],
             None, ['Known gap — production roadmap item']),
        ]),

        ('💼 Part 7 — Market & Business Questions', 'Market size, competitors, revenue model, data moat, next steps.', [
            (37, 'Market', 'What is the market size?',
             ['Global e-learning market: $400B by 2026. Online skills training: $50B+ growing 15% YoY. Workforce reskilling (B2B): 25% YoY driven by AI disruption. "Agentic personalized" segment: early-stage — we are entering at the right time.'],
             None, []),
            (38, 'Market', 'Who are competitors and how do you win?',
             ['Coursera/Udemy: static content, no adaptation → Win: dynamic agents vs static.',
              'LinkedIn Learning: no diagnostic, no adaptive plans → Win: real gap analysis.',
              'Duolingo (for tech): gamified but shallow → Win: professional depth.',
              'ChatGPT prompting: requires expert prompting, no memory → Win: autonomous orchestration.'],
             None, []),
            (39, 'Market', 'What is the business model?',
             ['<b>B2C:</b> Free (5 sessions) → Pro $19/mo → Lifetime $199.',
              '<b>B2B:</b> Enterprise workforce upskilling $99/user/year. Custom domain packs.',
              '<b>B2B2C:</b> White-label for universities/bootcamps. Revenue share on completion.',
              'Target: $50/user/year × 10,000 users = $500k ARR Year 1.'],
             None, []),
            (40, 'Market', 'How does your system quantify business impact?',
             ['SimulationAgent provides measurable projections: salary increase (e.g., +27% after Python), hiring time reduction (4-6 months → 2-3 months), job opening increase (+133%). Tangible ROI before investing a single hour studying.'],
             '/simulation — show salary impact metrics', []),
            (41, 'Market', 'What is the data moat?',
             ['As users complete sessions, we accumulate: quiz accuracy data (which questions predict success), skill gap patterns (common failure points per domain), learning velocity benchmarks. Over 100k users = proprietary dataset for fine-tuning — compounding moat generic AI cannot replicate.'],
             None, []),
            (42, 'Market', 'What data sources for market intelligence?',
             ['Currently: curated MARKET_SNAPSHOTS in MarketAgent.js — manually researched per domain. LLM-generated projections for any domain. Production roadmap: LinkedIn Jobs API (live job counts), Glassdoor/Levels.fyi (real-time salaries), GitHub Trending (rising skill detection).'],
             None, ['server/agent/MarketAgent.js']),
            (43, 'Market', 'What would you build next?',
             ['Priority: (1) LinkedIn Jobs API — show actual listings matching current skills. (2) Peer benchmarking — anonymized comparison vs same-goal learners. (3) Verifiable certificates — blockchain-backed skill proof. (4) Voice interface — speech recognition for practice sessions. (5) Mobile app — React Native, 30-min daily mode.'],
             None, []),
        ]),

        ('🔮 Part 8 — Innovation Questions', 'Frontier features, design innovations, learning science foundations.', [
            (44, 'Innovation', 'What is the single most innovative thing?',
             ['The <b>Agent Debate Engine</b>. When the system faces a decision, it pits 3 agents with different biases against each other: optimistic Advocate, cautious Critic, data-driven Analyst. They argue, vote with confidence weights, weighted majority wins. Visible to user — AI decision-making transparent and trustworthy.'],
             '/explain — Agent Debate History', ['server/agent/AgentDebate.js']),
            (45, 'Innovation', 'How does the Career Digital Twin differentiate you?',
             ["Digital Twins are established in manufacturing (GE, Siemens) but rare in career intelligence. Ours updates after each session: competency radar, mastery matrix, trajectory forecast, opportunity radar. It's a <b>living model that evolves</b>. Users watch their twin progress from Initializing → Expert."],
             '/career-twin', []),
            (46, 'Innovation', 'How does the What-If Simulator change career decision-making?',
             ['Traditional decisions are made with poor information. "Should I learn Kubernetes or React?" → our simulator makes it quantitative: +22% salary impact vs +15%, specific demand scores, Agent Verdict confidence. Turns vague choice into <b>data-driven investment decision</b>.'],
             '/simulation', []),
            (47, 'Innovation', 'What is the "Learn-First" session design innovation?',
             ['Traditional systems test before teaching. We invert: concept prime (definition + 4 points + example + pro tip) before every challenge. Grounded in spaced learning theory. Confidence calibration before/after measures metacognitive accuracy — a key self-directed learning skill.'],
             'Navigate to a session day', []),
            (48, 'Innovation', 'How does Skill Drift Detection address learning science?',
             ['Ebbinghaus forgetting curve (1885): skills not reinforced decay rapidly. SkillDriftAgent detects algorithmically — early vs recent scores on same skill. Drop &gt;20pts = superficial memorization. Logs intervention, recommends spaced repetition. Evidence-based adaptive learning, not just gamification.'],
             None, ['SmartAgent.js _detectSkillDrift()']),
        ]),
    ]

    for sec_title, sec_sub, questions in qa_sections:
        story.append(PageBreak())
        story.append(Paragraph(sec_title, SEC_HEAD))
        story.append(Paragraph(sec_sub, SEC_SUB))

        for num, cat, q, answer_lines, screen, proofs in questions:
            cat_colors = {
                'Technical': INDIGO, 'Architecture': VIOLET, 'Agentic': CYAN,
                'Security': RED, 'Market': EMERALD, 'Innovation': AMBER,
            }
            color = cat_colors.get(cat, INDIGO)
            story.append(qa_block(num, cat, q, answer_lines, screen, proofs, color))

        story.append(HR())

    # ── REQUIREMENT MATRIX ────────────────────────────────────────────────────
    story.append(Paragraph('✅ Part 9 — Requirement Coverage Matrix', SEC_HEAD))

    req_headers = ['Requirement', 'Status', 'File/Module', 'UI Location']
    req_rows = [
        ['Multi-agent architecture',     '✅', 'SmartAgent.js + 8 agent files',                '/demo, /explain'],
        ['Autonomous decision making',   '✅', 'Adaptor.js, _detectSkillDrift',                '/explain'],
        ['Agent Debate Engine',          '✅', 'AgentDebate.js',                               '/explain — Debate History'],
        ['Human-in-the-Loop',            '✅', 'Session.jsx confidence, Profiling.jsx',        '/session/:day, /profiling'],
        ['Short-term memory',            '✅', 'Session object in SmartAgent.js',              'All pages'],
        ['Long-term memory',             '✅', 'server/data/[userId].json',                    '/dashboard, /report'],
        ['Multi-step reasoning visible', '✅', 'agentDecisions array in session',              '/explain'],
        ['LLM integration',              '✅', 'GeminiService.js (Gemini + Groq)',             'All AI features'],
        ['LLM fallback (Groq)',          '✅', 'GeminiService.js _groqJSON/_groqText',         'Transparent'],
        ['Rule-based fallback',          '✅', 'RuleBase.js, QuizGenerator, Evaluator',        'Works offline'],
        ['Skill gap analysis',           '✅', 'Evaluator.scoreDiagnostic, MarketAgent',       '/diagnostic, /career-twin'],
        ['Personalized learning plan',   '✅', 'PlanBuilder.js',                              '/dashboard'],
        ['Adaptive plan modification',   '✅', 'Adaptor.js',                                  '/dashboard — addedByAgent'],
        ['Skill Drift Detection',        '✅', 'SmartAgent._detectSkillDrift',                '/explain'],
        ['Confidence calibration',       '✅', 'Session.jsx confidence phase',                '/session/:day'],
        ['Learn-first session design',   '✅', 'Session.jsx LearnPhase',                      '/session/:day'],
        ['AI-powered evaluation',        '✅', 'Evaluator.scoreSession (Gemini)',              '/session/:day result'],
        ['Report generation',            '✅', 'ReportGenerator.js',                          '/report'],
        ['Domain-specific quizzes',      '✅', 'RuleBase.js Medicine/Law/Civil/Finance',       '/diagnostic'],
        ['Career Digital Twin',          '✅', 'CareerTwin.jsx + forecastTrajectory',          '/career-twin'],
        ['What-If Simulator',            '✅', 'SimulationAgent.js + SimulationLab.jsx',       '/simulation'],
        ['Market Intelligence',          '✅', 'MarketAgent.js',                              '/career-twin'],
        ['Opportunity Radar',            '✅', 'SimulationAgent._buildOpportunityRadar',       '/career-twin'],
        ['Future Skill Forecast',        '✅', 'SimulationAgent.forecastTrajectory',           '/career-twin'],
        ['Explainability Console',       '✅', 'ExplainabilityConsole.jsx',                   '/explain'],
        ['Live Demo Mode (SSE)',          '✅', 'demo.js (SSE) + DemoMode.jsx',                '/demo'],
        ['Input validation',             '✅', 'goal.js, session.js routes',                  'Server-side'],
        ['Error boundaries',             '✅', 'ErrorBoundary.jsx',                           'App-level'],
        ['Lazy loading',                 '✅', 'App.jsx React.lazy per page',                 'App routing'],
        ['Responsive design',            '✅', 'All pages — Tailwind responsive',             'All pages'],
        ['Scalable architecture',        '✅', 'Modular agents, swappable storage',           'Architecture'],
        ['LLM resilience (3 tiers)',      '✅', 'Gemini → Groq → Rule-based → Static',        'Transparent'],
    ]

    req_data = [[Paragraph(h, TABLE_HEADER_STYLE) for h in req_headers]]
    for r in req_rows:
        req_data.append([
            Paragraph(r[0], TABLE_CELL_STYLE),
            Paragraph(f'<font color="#059669"><b>{r[1]}</b></font>', TABLE_CELL_STYLE),
            Paragraph(f'<font color="#4F46E5">{r[2]}</font>', TABLE_CELL_STYLE),
            Paragraph(r[3], TABLE_CELL_STYLE),
        ])

    req_table = Table(req_data, colWidths=[W*0.30, W*0.08, W*0.34, W*0.28])
    req_table.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,0),  SLATE900),
        ('ROWBACKGROUNDS',(0,1), (-1,-1), [WHITE, CARD_BG]),
        ('BOX',           (0,0), (-1,-1), 0.5, BORDER),
        ('INNERGRID',     (0,0), (-1,-1), 0.3, BORDER),
        ('TOPPADDING',    (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING',   (0,0), (-1,-1), 7),
        ('RIGHTPADDING',  (0,0), (-1,-1), 7),
        ('VALIGN',        (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(req_table)
    story.append(Spacer(1, 12))

    # Final score box
    final = Table([[Paragraph(
        '<b>Final Score: 32/32 requirements implemented · 9 specialized agents · 4 frontier pages · '
        'Full LLM resilience · Zero black-box decisions</b>',
        S('Final', fontSize=11, fontName='Helvetica-Bold', textColor=colors.HexColor('#065F46'), alignment=TA_CENTER))
    ]], colWidths=[W])
    final.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,-1), colors.HexColor('#ECFDF5')),
        ('BOX',           (0,0), (-1,-1), 1.5, EMERALD),
        ('TOPPADDING',    (0,0), (-1,-1), 12),
        ('BOTTOMPADDING', (0,0), (-1,-1), 12),
        ('LEFTPADDING',   (0,0), (-1,-1), 16),
        ('RIGHTPADDING',  (0,0), (-1,-1), 16),
    ]))
    story.append(final)
    story.append(Spacer(1, 6*mm))

    footer = Paragraph(
        'SKILL FORGE — Autonomous Multi-Agent Career Intelligence Platform · Hackathon 2026 · Confidential',
        S('Footer', fontSize=8, fontName='Helvetica', textColor=colors.HexColor('#94A3B8'), alignment=TA_CENTER)
    )
    story.append(footer)

    doc.build(story)
    print(f'[OK] PDF written to {path}')


# ── DOCX ──────────────────────────────────────────────────────────────────────
def build_docx(path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

    def set_font(run, bold=False, italic=False, size=11, color=None):
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)

    def add_heading(text, level=1, color=(79, 70, 229)):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
        return h

    def add_para(text, bold=False, italic=False, size=11, color=None, indent=0):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic, size=size, color=color)
        return p

    def shade_para(para, hex_color='F1F5F9'):
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        pPr.append(shd)

    # ── Cover ────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('SKILL FORGE')
    set_font(r, bold=True, size=32, color=(79, 70, 229))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Autonomous Multi-Agent Career Intelligence Platform')
    set_font(r, size=14, color=(71, 85, 105))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run('Judge Q&A Reference + Live Demo Guide  ·  Hackathon 2026  ·  Confidential')
    set_font(r, size=10, italic=True, color=(100, 116, 139))

    doc.add_paragraph()

    # Stats table
    t = doc.add_table(rows=2, cols=4)
    t.style = 'Table Grid'
    stats = [['50+', '9', '7', '100%'], ['Judge Q&As', 'Agent Types', 'Demo Steps', 'Requirements Met']]
    for i, row in enumerate(stats):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = val
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            if i == 0:
                set_font(run, bold=True, size=18, color=(79, 70, 229))
            else:
                set_font(run, size=9, color=(100, 116, 139))
    doc.add_paragraph()

    # ── Demo Guide ──────────────────────────────────────────────────────────
    add_heading('Part 1 — Live Demo Guide (Run This First)', level=1)

    tip = doc.add_paragraph()
    shade_para(tip, 'ECFDF5')
    r = tip.add_run('START HERE: Go to http://localhost:5173/demo → Select a goal → Click "Run Live Autonomous Career Analysis". '
                    'Watch 7 agents activate with real API calls in ~15 seconds.')
    set_font(r, size=10.5, color=(6, 95, 70))

    demo_steps_docx = [
        ('1', 'Open Landing Page', 'Show hero, 7-agent badge, Live Demo button, Frontier Modules grid.'),
        ('2', 'Select "Medical Doctor"', 'Professional domain — proves domain-specific intelligence.'),
        ('3', 'Click "Run Live Autonomous Career Analysis"', '7 agents activate sequentially with real API calls. SSE stream visible.'),
        ('4', 'Review Completion Summary', 'Skills mapped, sessions planned, demand score, salary projection.'),
        ('5', 'Open Career Digital Twin (/career-twin)', 'Competency radar, 12-month trajectory, market intelligence panel.'),
        ('6', 'Simulation Lab → "Machine Learning"', '+30% salary impact, demand jump 58→88, Agent Verdict confidence score.'),
        ('7', 'Explainability Console (/explain)', 'Every agent decision with reasoning, confidence bars, debate transcripts.'),
    ]
    for num, title_s, body_s in demo_steps_docx:
        p = doc.add_paragraph(style='List Number')
        r = p.add_run(f'Step {num}: {title_s}  ')
        set_font(r, bold=True, size=11, color=(15, 23, 42))
        r2 = p.add_run(body_s)
        set_font(r2, size=10.5, color=(55, 65, 81))

    doc.add_paragraph()

    # ── Q&A sections ─────────────────────────────────────────────────────────
    qa_data_docx = [
        ('Technical Questions', [
            (1,  'What is the tech stack?',
             'Frontend: React 18 + Vite + TailwindCSS. Backend: Node.js + Express. LLMs: Gemini 2.0 Flash + Groq llama-3.3-70b (fallback). Storage: JSON (upgrade-ready to PostgreSQL).'),
            (2,  'What happens if Gemini quota is exceeded?',
             'GeminiService.js has automatic Groq fallback. HTTP 429 → immediately switches to Groq (free, llama-3.3-70b). If both fail → rule-based logic. System never crashes.'),
            (3,  'Are diagnostic questions domain-specific?',
             'Yes. Three tiers: (1) RuleBase — real MCQs for Medicine/Law/Civil/Finance. (2) Gemini — for any other domain. (3) Knowledge-testing fallback templates. Never meta-questions.'),
            (4,  'How does evaluation work?',
             'Three-tier scoring: Gemini AI analysis → keyword/depth scoring → normalized grade 0-100. Confidence calibration compares predicted vs actual score.'),
            (5,  'How is the plan personalized?',
             'PlanBuilder applies proficiency multiplier: <60% = more days, >80% = compressed. Sorted weakest-first. AdaptorAgent re-evaluates every 3 sessions.'),
            (6,  'How does the What-If Simulator work?',
             'SimulationAgent sends goal/skills/domain + proposed skill to LLM. Gets salary delta, demand shift, hiring time improvement, timeline, alternatives, risk factors.'),
            (7,  'How does the Live Demo work technically?',
             'Server-Sent Events (SSE) — real-time server→browser stream. Real API calls, not mocked. React EventSource API animates each agent as server processes.'),
            (8,  'How do you prevent React UI rollback?',
             'Cancelled flag cleanup pattern: let cancelled = false, set true on cleanup. API callbacks check !cancelled before setState. Prevents React 18 Strict Mode double-invoke.'),
        ]),
        ('Architecture Questions', [
            (9,  'Walk through the full agent pipeline.',
             '1.GoalAgent → 2.DecomposeAgent → 3.DiagnosticAgent → 4.ScoringAgent → 5.CurriculumAgent → 6.EvaluatorAgent → 7.AdaptorAgent → 8.MarketAgent → 9.SimulationAgent.'),
            (10, 'How is state shared between agents?',
             'Blackboard pattern: all agents receive the session object. SmartAgent.js is the orchestrator — passes session to each agent, merges results. No direct agent-to-agent dependencies.'),
            (11, 'Is the system truly agentic?',
             'Yes. AdaptorAgent decides based on data (not LLM). AgentDebate has 3 biased agents arguing. SkillDriftAgent self-initiates alerts. Rule-base bypasses LLM when available.'),
            (12, 'How would you scale to millions of users?',
             'PostgreSQL + JSONB. BullMQ/Redis for LLM queuing. Redis cache for domain data. Express is stateless — loadbalancer-ready. Concurrent agent steps via Promise.all.'),
        ]),
        ('Agentic Reasoning Questions', [
            (13, 'Explain the Agent Debate mechanism.',
             '3 agents with biases: AdvocateAgent (optimistic), CriticAgent (cautious), AnalystAgent (data-driven). Each votes with confidence weight. Weighted majority wins. Visible in Explainability Console.'),
            (14, 'Is there Human-in-the-Loop?',
             'Yes: confidence self-rating, profiling questionnaire, diagnostic quiz, session override (navigate to any day), Simulation Lab queries.'),
            (15, 'What is the Career Digital Twin?',
             'A living virtual model of the user: competency radar, mastery trajectory, market position, opportunity gaps. Updates after every session. Level: Initializing → Expert.'),
            (16, 'What is the single most innovative thing?',
             'The Agent Debate Engine. 3 agents with different biases argue before every plan change. Produces calibrated decisions visible to the user — zero black-box behavior.'),
        ]),
        ('Security & Market Questions', [
            (17, 'How are API keys secured?',
             'Stored in .env (gitignored). Read fresh at runtime via getter — never cached. Production: AWS/GCP Secrets Manager.'),
            (18, 'What is the business model?',
             'B2C: Free → Pro $19/mo → Lifetime $199. B2B: $99/user/year enterprise upskilling. B2B2C: white-label for universities. Target: $500k ARR Year 1 at 10k users.'),
            (19, 'Who are competitors and how do you win?',
             'Coursera/Udemy: static content → Win: dynamic agents. LinkedIn Learning: no diagnostic → Win: real gap analysis. ChatGPT: no memory → Win: autonomous orchestration.'),
            (20, 'What is the data moat?',
             'Proprietary dataset: which questions predict success, common failure patterns, learning velocity benchmarks. Over 100k users = fine-tuning data competitors cannot replicate.'),
        ]),
    ]

    for section_name, questions_d in qa_data_docx:
        add_heading(section_name, level=2, color=(79, 70, 229))
        for num, q, a in questions_d:
            qp = doc.add_paragraph()
            shade_para(qp, 'F1F5F9')
            r = qp.add_run(f'Q{num}. {q}')
            set_font(r, bold=True, size=11.5, color=(15, 23, 42))

            ap = doc.add_paragraph()
            shade_para(ap, 'F8FAFC')
            r = ap.add_run(a)
            set_font(r, size=10.5, color=(55, 65, 81))
            doc.add_paragraph()

    # ── Coverage table ───────────────────────────────────────────────────────
    add_heading('Requirement Coverage Matrix', level=2, color=(5, 150, 105))

    reqs_docx = [
        'Multi-agent architecture (9 agents)',
        'Autonomous decision making (Adaptor, SkillDrift)',
        'Agent Debate Engine (3-agent voting)',
        'Human-in-the-Loop (confidence, profiling)',
        'Short-term + long-term memory',
        'Multi-step reasoning (visible log)',
        'LLM integration (Gemini 2.0 Flash)',
        'LLM fallback (Groq, free tier)',
        'Rule-based fallback (works offline)',
        'Skill gap analysis',
        'Personalized learning plan',
        'Adaptive plan modification',
        'Skill Drift Detection',
        'Confidence calibration',
        'Learn-first session design',
        'AI-powered evaluation',
        'Report generation',
        'Domain-specific quizzes (Medicine/Law/Civil/Finance)',
        'Career Digital Twin',
        'What-If Simulator',
        'Market Intelligence Agent',
        'Opportunity Radar',
        'Future Skill Forecast',
        'Explainability Console',
        'Live Demo Mode (SSE)',
        'Input validation + error handling',
        'Responsive design',
        'Scalable modular architecture',
    ]

    t2 = doc.add_table(rows=len(reqs_docx)+1, cols=2)
    t2.style = 'Table Grid'
    for j, h in enumerate(['Requirement', 'Status']):
        cell = t2.cell(0, j)
        cell.text = h
        p = cell.paragraphs[0]
        run = p.runs[0]
        set_font(run, bold=True, size=10, color=(248, 250, 252))
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = _OE('w:shd')
        shd.set(_qn('w:val'), 'clear')
        shd.set(_qn('w:color'), 'auto')
        shd.set(_qn('w:fill'), '0F172A')
        tcPr.append(shd)

    for i, req in enumerate(reqs_docx):
        row = t2.rows[i+1]
        row.cells[0].text = req
        row.cells[1].text = '✅ Implemented'
        set_font(row.cells[0].paragraphs[0].runs[0], size=9.5)
        r = row.cells[1].paragraphs[0].runs[0]
        set_font(r, bold=True, size=9.5, color=(5, 150, 105))

    doc.add_paragraph()
    final_p = doc.add_paragraph()
    shade_para(final_p, 'ECFDF5')
    r = final_p.add_run(f'Final Score: {len(reqs_docx)}/28+ requirements implemented · 9 specialized agents · 4 frontier pages · '
                         'Full LLM resilience · Zero black-box decisions')
    set_font(r, bold=True, size=11, color=(6, 95, 70))

    doc.save(path)
    print(f'[OK] DOCX written to {path}')


if __name__ == '__main__':
    import sys, os
    base = r'C:\CODING\HACKap'
    pdf_path  = os.path.join(base, 'judge-qa.pdf')
    docx_path = os.path.join(base, 'judge-qa.docx')

    print('Building PDF...')
    build_pdf(pdf_path)

    print('Building DOCX...')
    try:
        build_docx(docx_path)
    except ImportError:
        print('python-docx not installed — skipping DOCX. Run: pip install python-docx')

    print('\nDone!')
